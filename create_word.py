from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys

# ── helpers ────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def cell_text(cell, text, bold=False, size=10, color=None,
              align=WD_ALIGN_PARAGRAPH.CENTER, rtl=True):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    # RTL paragraph
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1' if rtl else '0')
    pPr.append(bidi)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

# ── data ───────────────────────────────────────────────────────────────────
TASKS = [
    ('מקרר ביצים',               'מחסן -1',      'ראשון'),
    ('מקרר חלב + מדפים',         'מחסן -1',      'ראשון'),
    ('מקרר ירקות',               'מחסן -1',      'שני'),
    ('מקפיא קפואים',             'מחסן -1',      'שני'),
    ('מקפיא מאפית',              'מאפית',        'שלישי'),
    ('מקרר מאפית',               'מאפית',        'שלישי'),
    ('מקרר בשר',                 'קצביה',        'רביעי'),
    ('מקרר עופות',               'קצביה',        'רביעי'),
    ('גיבוי גבינות',             'מעדניה גבוי',  'חמישי'),
    ('מקפיא מאפית 2',            'מחסן -1',      'חמישי'),
    ('מקרר החזרות',              'מחסן -1',      'חמישי'),
    ('ניקיון וסדר רמפה',         'רמפה',         'שישי'),
    ('מקרר דגים ועופות חיצוני', 'אולם',         'שישי'),
]

DAYS_ORDER = ['ראשון','שני','שלישי','רביעי','חמישי','שישי']

DAY_COLORS = {
    'ראשון':  'FFF9C4',
    'שני':    'E8F5E9',
    'שלישי':  'E3F2FD',
    'רביעי':  'FBE9E7',
    'חמישי':  'F3E5F5',
    'שישי':   'FFF3E0',
}

YELLOW = '1A1A2E'   # header bg
HEADER_TEXT = 'FFCC00'

# ── document ───────────────────────────────────────────────────────────────
doc = Document()

# Page: A4 landscape
section = doc.sections[0]
section.page_width  = Cm(29.7)
section.page_height = Cm(21.0)
section.left_margin  = Cm(1.5)
section.right_margin = Cm(1.5)
section.top_margin   = Cm(1.5)
section.bottom_margin= Cm(1.5)

# RTL section
sectPr = section._sectPr
bidi = OxmlElement('w:bidi')
sectPr.append(bidi)

# ── TITLE ──────────────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_p.add_run('לוז ניקיון מקררים ומקפיאים — שופרסל')
tr.bold = True
tr.font.size = Pt(18)
tr.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub_p.add_run('גרף שבועי | עדכון בתחילת כל שבוע | יוני 2026')
sr.font.size = Pt(11)
sr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_paragraph()

# ── MAIN TABLE ─────────────────────────────────────────────────────────────
col_headers = ['#', 'שם המקרר / מקפיא', 'מיקום', 'יום ניקיון',
               'ראשון', 'שני', 'שלישי', 'רביעי', 'חמישי', 'שישי']

table = doc.add_table(rows=1+len(TASKS), cols=len(col_headers))
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
hrow = table.rows[0]
for ci, h in enumerate(col_headers):
    c = hrow.cells[ci]
    set_cell_bg(c, YELLOW)
    cell_text(c, h, bold=True, size=10, color=HEADER_TEXT)

# Data rows
for ri, (name, loc, clean_day) in enumerate(TASKS, start=1):
    row = table.rows[ri]
    bg = DAY_COLORS.get(clean_day, 'FFFFFF')

    # #
    set_cell_bg(row.cells[0], bg)
    cell_text(row.cells[0], str(ri), size=9)
    # Name
    set_cell_bg(row.cells[1], bg)
    cell_text(row.cells[1], name, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
    # Location
    set_cell_bg(row.cells[2], bg)
    cell_text(row.cells[2], loc, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    # Day
    set_cell_bg(row.cells[3], bg)
    cell_text(row.cells[3], clean_day, bold=True, size=10, color='1A1A2E')
    # Day checkboxes (4-9)
    for di, d in enumerate(DAYS_ORDER):
        ci = 4 + di
        set_cell_bg(row.cells[ci], 'FFFFFF')
        if d == clean_day:
            set_cell_bg(row.cells[ci], bg)
            cell_text(row.cells[ci], '☐', size=14, color='1A1A2E')
        else:
            cell_text(row.cells[ci], '', size=9)

# Column widths
col_widths = [0.7, 5.5, 2.5, 2.0, 2.0, 2.0, 2.0, 2.0, 2.0, 2.0]
for ci, w in enumerate(col_widths):
    for row in table.rows:
        row.cells[ci].width = Cm(w)

# Row heights
for row in table.rows:
    row.height = Cm(0.9)

doc.add_paragraph()

# ── PERIODIC TABLE ─────────────────────────────────────────────────────────
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r2 = p2.add_run('📋  משימות תקופתיות (עומק)')
r2.bold = True
r2.font.size = Pt(13)
r2.font.color.rgb = RGBColor(0xE6, 0x7E, 0x22)

peri_table = doc.add_table(rows=3, cols=4)
peri_table.style = 'Table Grid'
peri_table.alignment = WD_TABLE_ALIGNMENT.CENTER

peri_headers = ['משימה', 'תאריך 1', 'תאריך 2', 'תאריך 3']
for ci, h in enumerate(peri_headers):
    c = peri_table.rows[0].cells[ci]
    set_cell_bg(c, YELLOW)
    cell_text(c, h, bold=True, size=10, color=HEADER_TEXT)

peri_data = [
    ('ניקיון מדף תחתון — מקרר ירקות',
     'שישי 19.6.2026', 'שישי 19.7.2026', 'שישי 19.8.2026'),
    ('ניקיון מדף תחתון — מקרר חלב',
     'שלישי-חמישי 23-25.6', 'שלישי-חמישי 23-25.8', 'שלישי-חמישי 23-25.10'),
]
for ri, row_data in enumerate(peri_data, start=1):
    for ci, val in enumerate(row_data):
        c = peri_table.rows[ri].cells[ci]
        set_cell_bg(c, 'FFF8E1')
        align = WD_ALIGN_PARAGRAPH.RIGHT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
        cell_text(c, val, size=10, bold=(ci==0), align=align)

peri_table.rows[0].height = Cm(0.85)
for ri in range(1,3):
    peri_table.rows[ri].height = Cm(0.85)

doc.add_paragraph()

# ── LEGEND ─────────────────────────────────────────────────────────────────
leg = doc.add_paragraph()
leg.alignment = WD_ALIGN_PARAGRAPH.RIGHT
lr = leg.add_run('☐ = טרם בוצע  |  ☑ = בוצע  |  ❌ = לא בוצע (סמן בעט אדום)')
lr.font.size = Pt(9)
lr.font.color.rgb = RGBColor(0x77,0x77,0x77)

# Footer
doc.add_paragraph()
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run('פותח עבור שופרסל על ידי Evgeny Tsipis | Claude AI · יוני 2026')
fr.font.size = Pt(8)
fr.font.color.rgb = RGBColor(0xAA,0xAA,0xAA)

# ── SAVE ───────────────────────────────────────────────────────────────────
out = r'C:\Users\etsip\OneDrive\TSIPIS\Job\שופרסל\לוז ניקיון מקררים ומקפיאים - עדכון.docx'
doc.save(out)
sys.stdout.buffer.write(b'Word saved OK\n')
